import cx_Oracle
from docx import Document

username = 'CS_COURSE'
password = 'Cs671bsas_3WE'
host = '218.7.43.141'
port = '1521'
service_name = 'CS_COURSE.JSKB'

dsn_tns = cx_Oracle.makedsn('218.7.43.141', '1521', service_name='orcl') 
connection = cx_Oracle.connect(user='CS_COURSE', password='Cs671bsas_3WE', dsn=dsn_tns)

cursor = connection.cursor()
cursor_lend = connection.cursor()
query = "SELECT KCM,ZC,JASDM,SKXQ,KSJC,JSJC,XNXQDM,JASMC,JSM,XS,JASDM FROM CS_COURSE.JSKB"
#                0  1    2    3    4    5     6      7    8  9    10
query_lend = "SELECT JASDM,JASMC,XNXQDM,ZC,XQ,KSJC,JSJC,SHZT FROM CS_COURSE.CLASSROOM_LEND"
#                      0     1     2    3  4   5    6    7
cursor.execute(query)
cursor_lend.execute(query_lend)

results = cursor.fetchall()
results_lend = cursor_lend.fetchall()

cursor.close()
cursor_lend.close()
connection.close()

#确定年份
def loc_year(year_string):
    x = year_string.split('-')
    return x


#确定教室
def loc_classroom(room_string):
    if room_string == None:
        return False
    if room_string[:4] == '计算中心':
        return True
    return False


#确定第几周
def loc_week(week_string):
    week_list = list()
    weeks = week_string.replace("周", "").split(",")
    for week in weeks:
        if "-" in week:
            start, end = week.split("-")
            for i in range(int(start), int(end)+1):
                week_list.append(i)
        else:
            week_list.append(int(week))
    
    return week_list


#确定周几
weekdays_table={'星期一':1,
                '星期二':2,
                '星期三':3,
                '星期四':4,
                '星期五':5,
                '星期六':6,
                '星期天':7}
def loc_weekday(weekday_string):
    return weekdays_table[weekday_string]


#确定第几节
def loc_class_time(start,end):
    time=list()
    temp = 0
    if start == 1:time.append(1)
    elif start == 3:time.append(2)
    else:time.append(int(start/2))
    flag = int(end)
    while flag-int(start)>=3:
        time.append(time[-1]+1)
        flag -= 3
    return time

temp_year=input()
temp_week=input()
processed_year=loc_year(temp_year)

#表格
from docx.oxml.ns import qn
from docx.shared import Pt,Cm,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(5)

#获取本文档中的所有章节
sections = doc.sections
#将该章节中的纸张方向设置为横向
for section in sections:
    #需要同时设置width,height才能成功
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width 
    section.page_height = new_height

headl = doc.add_heading("哈尔滨工程大学 教室课表")
headl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
head2 = doc.add_heading('',level=1).add_run(f'学年学期：{processed_year[0]}-{processed_year[1]}学年第{processed_year[2]}学期  第{temp_week}周  所属功能区：计算机房  所属教学楼：机房  座位数：64')
head2.font.size = Pt(9)
head2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table = doc.add_table(rows=6,cols=8,style='Table Grid')
for row in table.rows:row.height = Pt(25)
for col in table.columns:col.width = Pt(100)
table.rows[0].height = Pt(25)
table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

header_cells=table.rows[0].cells
header_cells[1].text = '星期一'
header_cells[2].text = '星期二'
header_cells[3].text = '星期三'
header_cells[4].text = '星期四'
header_cells[5].text = '星期五'
header_cells[6].text = '星期六'
header_cells[7].text = '星期日'
table.cell(1,0).text = '第一大节'
table.cell(2,0).text = '第二大节'
table.cell(3,0).text = '第三大节'
table.cell(4,0).text = '第四大节'
table.cell(5,0).text = '第五大节'



#遍历数据
for mesg in results:
    if temp_year != mesg[6]:continue
    if loc_classroom(mesg[7]) == False:continue
    this_weekday = loc_weekday(mesg[3])
    this_time = loc_class_time(int(mesg[4]),int(mesg[5]))
    this_week = loc_week(mesg[1])
    if this_week.count(int(temp_week)):
        print('you')
        for time in this_time:
            mesg_str = f'(本){mesg[0]},{mesg[8]},{mesg[9]}人\n{mesg[1]},{mesg[3]},{mesg[4]}-{mesg[5]}节,{mesg[7]}'
            table.cell(time,this_weekday).add_paragraph(mesg_str)
            print(mesg_str)

for mesg in results_lend:
    if temp_year != mesg[2]:continue
    this_weekday = int(mesg[4])
    this_time = loc_class_time(int(mesg[5]),int(mesg[6]))
    this_week = loc_week(mesg[3])
    if this_week.count(int(temp_week)):
        print('zai')
        for time in this_time:
            mesg_str = f'{mesg[1]},{mesg[2]}借用'
            table.cell(time,this_weekday).add_paragraph(mesg_str)
            print(mesg_str)

# 保存Word文档
doc.save('timetable-final.docx')

# 这是b
# 这是啥b
# push你妈